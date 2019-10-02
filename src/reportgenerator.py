"""Module that handles the creation of a docx report."""

# coding=utf-8

import re
import json

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

from conf.report import REPORT_TEMPLATE_DIR, REPORT_TEMPLATE_MAIN, REPORT_TEMPLATE_BASE, LANGUAGES
from src.cvss import vuln_risk_level, vuln_cvssv3


class Generator:
    """Defines the generation of a report."""
    styleRegex = re.compile(r"\[\[(B?I?U?)\]\](.*?)\[\[/B?I?U?\]\]")
    hyperlinkRegex = re.compile(
        r"\[\[HYPERLINK\]\](.*?)\|\|(.*?)\[\[/HYPERLINK\]\]")
    imageRegex = re.compile(r"\[\[IMAGE\]\](.*?)\|\|(.*?)\[\[/IMAGE\]\]")

    @staticmethod
    def __escape_str(string):
        """Escapres \a, \b, \f, \r, \t, \v"""
        return string.replace("\a", "\\a").replace("\b", "\\b").replace(
            "\f", "\\f").replace("\r", "\\r").replace("\t", "\\t").replace("\v", "\\v")

    @staticmethod
    def generate_report(json_content):
        """Main process of the report generation."""
        if isinstance(json_content, str):
            return json_content.replace("%", r"\%")

        if "name" not in json_content:
            json_content["name"] = ""
        if "content" not in json_content:
            json_content["content"] = ""

        file_content = None
        with open("templates/" + json_content["type"] + ".tex") as file_stream:
            file_content = file_stream.read()
        final_content = file_content
        final_content = re.sub(
            "##.*NAME##", json_content["name"], final_content)

        content = ""
        if isinstance(json_content["content"], list):
            for element in json_content["content"]:
                content += Generator.generate_report(element)
        elif isinstance(json_content["content"], dict):
            content += Generator.generate_report(json_content["content"])
        elif isinstance(json_content["content"], str):
            content += Generator.generate_report(json_content["content"])

        final_content = re.sub("##.*CONTENT##", content, final_content)
        return Generator.__escape_str(final_content)

    @staticmethod
    def __sub_dict(dic, text):
        for name, value in dic.items():
            if isinstance(value, str):
                text = re.sub("##" + name + "##", value, text)
                match = re.search("##" + name + r"#(\d+)##",
                                  text, re.MULTILINE)
                if match is not None:
                    text = re.sub("##" + name + "#" + match.group(1) +
                                  "##", value[int(match.group(1))], text)
                match = re.search("##" + name + r"#(\d+):(\d+)##",
                                  text, re.MULTILINE)
                if match is not None:
                    text = re.sub("##" + name + "#" + match.group(1) + ":" + match.group(2) + "##",
                                  value[int(match.group(1)):int(match.group(2))], text)
            elif isinstance(value,int) or isinstance(value,float):
                text = re.sub("##" + name + "##", str(value), text)
            else:
                pass #value should be a list or an orderedDict so nothing to do

        return text

    @staticmethod
    def __do_fill(dic, content):
        if isinstance(dic, str):
            dic = Generator.__sub_dict(content, dic)
            return dic

        if isinstance(dic, list):
            res = []
            for entry in dic:
                res.append(Generator.__do_fill(entry, content))
            return res

        if "type" not in dic:
            return dic

        if dic["type"].find("##") >= 0:
            dic["type"] = Generator.__sub_dict(content, dic["type"])

        if dic["type"] == "unordered_list" or dic["type"] == "ordered_list":
            res_content = []
            for entry in content[dic["filer"]].values():
                for content_var in dic["content"]:
                    template = dict(content_var)
                    template["content"] = Generator.__do_fill(
                        template["content"], entry)
                    res_content.append(template)
            dic["content"] = res_content
            return dic

        if dic["type"].startswith("Vulns-"):
            res_content = []
            match = dic["type"].split("-")[1]
            for doc_id, vuln in content["Vulns"].items():
                vuln["doc_id"] = doc_id
                if "riskLvl" not in vuln:
                    vuln["riskLvl"], vuln["impLvl"], vuln["expLvl"] = vuln_risk_level(
                        vuln)
                    vuln["cvss"], vuln["cvssImp"], vuln["cvssExp"] = vuln_cvssv3(
                        vuln)

                if (vuln["riskLvl"] == match and vuln["status"] == "Vulnerable") or (
                        vuln["status"] == match):
                    for content_var in dic["content"]:
                        template = dict(content_var)
                        template["content"] = Generator.__do_fill(
                            template["content"], vuln)
                        res_content.append(template)
            dic["content"] = res_content
            return dic

        if dic["type"] == "VulnsFull":
            l_res = []
            cat = None
            sub_cat = None
            for doc_id, vuln in content["Vulns"].items():
                vuln["doc_id"] = doc_id
                if vuln["category"] != cat:
                    cat = vuln["category"]
                    sub_cat = None
                    template = Generator.__do_fill(
                        dict(dic["catContent"]), vuln)
                    l_res.append(template)

                if vuln["sub_category"] != sub_cat:
                    sub_cat = vuln["sub_category"]
                    template = Generator.__do_fill(
                        dict(dic["subcatContent"]), vuln)
                    l_res.append(template)

                for content in dic["content-" + vuln["status"]]:
                    template = Generator.__do_fill(dict(content), vuln)
                    l_res.append(template)
            dic["content"] = l_res
            return dic

        if "filer" in dic:
            content = content[dic["filer"]]

        if "content" not in dic:
            return dic

        if isinstance(dic["content"], list):
            l_res = []
            for element in dic["content"]:
                l_res.append(Generator.__do_fill(element, content))
            dic["content"] = l_res
        elif isinstance(dic["content"], str):
            dic["content"] = Generator.__sub_dict(content, dic["content"])
        else:
            dic["content"] = Generator.__do_fill(dic["content"], content)

        return dic

    @staticmethod
    def generate_json(json_content, template):
        """Loads template file into json."""
        template_path = REPORT_TEMPLATE_DIR + template + "/"
        structure = None
        with open(template_path + REPORT_TEMPLATE_MAIN, "r") as file_stream:
            structure = file_stream.read()
            structure = json.loads(structure)

        result_json = {}
        result_json["type"] = "report"
        result_json["content"] = []
        for element in structure:
            with open(template_path + element + ".json", "r") as file_stream:
                file_content = file_stream.read()
            json_file_content = json.loads(file_content)
            res = Generator.__do_fill(json_file_content, json_content)
            result_json["content"].append(res)
        return result_json

    @staticmethod
    def __cut_before(string, match):
        return string[string.find(match) + len(match):]

    @staticmethod
    def __cut_after(string, match):
        return string[0:string.find(match)]

    @staticmethod
    def __get_body(xml_content):
        out = Generator.__cut_before(xml_content, "<w:body>")
        out = Generator.__cut_after(out, "</w:body>")
        return out

    @staticmethod
    def __align(alignment):
        return getattr(WD_PARAGRAPH_ALIGNMENT, alignment)

    @staticmethod
    def __v_align(alignment):
        return getattr(WD_CELL_VERTICAL_ALIGNMENT, alignment)

    @staticmethod
    def __generate_table(document, json_input, template):
        table = document.add_table(
            json_input["row"], json_input["col"], json_input["style"])
        for row in range(0, json_input["row"]):
            for col in range(0, json_input["col"]):
                Generator.generate_docx(table.cell(row, col),
                                        json_input["content"][row][col],
                                        template)
                if "alignment" in json_input:
                    for paragraph in table.cell(row, col).paragraphs:
                        paragraph.alignment = Generator.__align(
                            json_input["alignment"])

                if "colAlignment" in json_input:
                    for paragraph in table.cell(row, col).paragraphs:
                        paragraph.alignment = Generator.__align(
                            json_input["colAlignment"][col])
                if "rowAlignment" in json_input:
                    for paragraph in table.cell(row, col).paragraphs:
                        paragraph.alignment = Generator.__align(
                            json_input["rowAlignment"][row])

                if "celAlignment" in json_input:
                    for paragraph in table.cell(row, col).paragraphs:
                        paragraph.alignment = Generator.__align(
                            json_input["celAlignment"][row][col])

                if "vAlignment" in json_input:
                    table.cell(row, col).vertical_alignment = Generator.__v_align(
                        json_input["vAlignment"])

                if "colVAlignment" in json_input:
                    table.cell(row, col).vertical_alignment = Generator.__v_align(
                        json_input["colVAlignment"][col])

                if "rowVAlignment" in json_input:
                    table.cell(row, col).vertical_alignment = Generator.__v_align(
                        json_input["rowVAlignment"][row])

                if "celVAlignment" in json_input:
                    table.cell(row, col).vertical_alignment = Generator.__v_align(
                        json_input["celVAlignment"][row][col])

                if "width" in json_input:
                    col = 0
                    for width in json_input["width"]:
                        table.columns[col].width = Cm(width)
                        col += 1

                if "firstCol" in json_input:
                    table.first_col = json_input["firstCol"]
                if "firstRow" in json_input:
                    table.first_row = json_input["firstRow"]
                if "lastCol" in json_input:
                    table.last_col = json_input["lastCol"]
                if "lastRow" in json_input:
                    table.last_row = json_input["lastRow"]
                if "hBand" in json_input:
                    table.h_band(json_input["hBand"])
                if "vBand" in json_input:
                    table.v_band(json_input["vBand"])

    @staticmethod
    def generate_docx(document, json_input, template):
        """Creates a docx file from a template and a list of values in a json."""
        if isinstance(json_input, str):
            document.text = json_input
            return

        if "type" in json_input:
            if json_input["type"] == "image":
                width = None
                height = None
                if "width" in json_input:
                    width = Cm(json_input["width"])
                if "height" in json_input:
                    height = Cm(json_input["height"])
                path = json_input["path"]
                if path[0] != '/':
                    path = REPORT_TEMPLATE_DIR + template + "/" + path
                document.add_picture(path, width, height)
            if json_input["type"] == "table":
                Generator.__generate_table(document, json_input, template)

            if json_input["type"] == "document":
                new_doc = Document(REPORT_TEMPLATE_DIR +
                                   template + "/" + json_input["path"])

                for paragraph in new_doc.paragraphs:
                    text = paragraph.text
                    style = paragraph.style.name
                    paragraph = document.add_paragraph(text, style)
                    paragraph.paragraph_format = paragraph.paragraph_format

        if "name" in json_input:
            paragraph = document.add_paragraph(
                json_input["name"], json_input["type"])
            if "alignment" in json_input:
                paragraph.alignment = Generator.__align(
                    json_input["alignment"])

        if "content" in json_input:
            if isinstance(json_input["content"], str):
                if "type" not in json_input:
                    json_input["type"] = "Normal"

                paragraph = document.add_paragraph("", json_input["type"])

                if "alignment" in json_input:
                    paragraph.alignment = Generator.__align(
                        json_input["alignment"])

                hyperlink_split = Generator.hyperlinkRegex.split(
                    json_input["content"])
                for cpt, hyperlink_split_value in enumerate(hyperlink_split):
                    if cpt % 3 == 0:
                        txt1 = hyperlink_split_value
                        image_split = Generator.imageRegex.split(txt1)
                        for cpt2, image_split_value in enumerate(image_split):
                            if cpt2 % 3 == 0:
                                txt2 = image_split_value
                                style_split = Generator.styleRegex.split(txt2)
                                for cpt3, style_split_value in enumerate(style_split):
                                    if cpt3 % 3 == 0:
                                        paragraph.add_run(style_split_value)
                                    elif cpt3 % 3 == 1:
                                        bold = paragraph.style.font.bold
                                        italic = paragraph.style.font.italic
                                        underline = paragraph.style.font.underline
                                        if style_split_value.find("B") >= 0:
                                            bold = True
                                        if style_split_value.find("I") >= 0:
                                            italic = True
                                        if style_split_value.find("U") >= 0:
                                            underline = True
                                    else:
                                        run = paragraph.add_run(
                                            style_split_value)
                                        run.bold = bold
                                        run.italic = italic
                                        run.underline = underline
                            elif cpt2 % 3 == 1:
                                run = paragraph.add_run()
                                path = image_split_value
                                if path[0] != '/':
                                    path = REPORT_TEMPLATE_DIR + template + "/" + path
                                run.add_picture(
                                    path, Cm(int(image_split[cpt2+1])))
                            else:
                                continue
                    elif cpt % 3 == 1:
                        paragraph.add_hyperlink(
                            hyperlink_split_value, hyperlink_split[cpt+1], style="Hyperlink")
                    else:
                        continue

            elif isinstance(json_input["content"], list):
                for content in json_input["content"]:
                    Generator.generate_docx(document, content, template)
            else:
                Generator.generate_docx(
                    document, json_input["content"], template)

    @staticmethod
    def generate_all(values, output_filename):
        """Main process of report generation."""
        template = values["Mission"]["template"]

        if len(LANGUAGES) > 1:
            lang = values["Mission"]["language"]
            if lang != LANGUAGES[0]:
                for vuln in values["Vulns"].values():
                    keys = list(vuln.keys()) # copy to avoid iterate on new items
                    for name in keys:
                        if name.find(lang) == len(name)-len(lang):
                            vuln[name[:-len(lang)]] = vuln[name]

        json_values = Generator.generate_json(values, template)

        doc = Document(docx=REPORT_TEMPLATE_DIR +
                       template + "/" + REPORT_TEMPLATE_BASE)
        Generator.generate_docx(doc, json_values, template)
        doc.save(output_filename)
